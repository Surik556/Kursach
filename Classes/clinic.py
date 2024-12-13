# clinic.py

import os
from datetime import datetime
from docx import Document
from appointment import Appointment

class Clinic:
    def __init__(self):
        self.appointments = []  # Список записей на прием

    def add_appointment(self, appointment: Appointment):
        self.appointments.append(appointment)  # Добавляем запись в список

    def save_to_docx(self, filename: str, veterinarian_name: str):
        document = Document()  # Создаем новый документ
        document.add_heading('Записи на прием', level=1)  # Заголовок документа

        for appt in self.appointments:
            document.add_heading(f'Запись на прием для {appt.animal.name}', level=2)
            document.add_paragraph(f'Клиент: {appt.client.name}')
            document.add_paragraph(f'Телефон: {appt.client.phone}')
            document.add_paragraph(f'Email: {appt.client.email}')
            document.add_paragraph(f'Дата приема: {appt.date}')
            document.add_paragraph(f'Время приема: {appt.time}')
            document.add_paragraph('---')  # Разделитель между записями

        directory = 'appointment'
        if not os.path.exists(directory):
            os.makedirs(directory)

        current_time = datetime.now().strftime("%H-%M-%S")
        file_path = f'{directory}/{veterinarian_name}_{filename}_{current_time}.docx'
        document.save(file_path)  # Сохраняем документ

        self.save_to_vet_file(veterinarian_name, filename, current_time)

        return file_path  # Возвращаем путь к сохраненному файлу

    def save_to_vet_file(self, veterinarian_name: str, filename: str, current_time: str):
        vet_directory = 'appointment_for_vet'
        if not os.path.exists(vet_directory):
            os.makedirs(vet_directory)

        vet_document = Document()  # Создаем новый документ для ветеринара
        vet_document.add_heading('Записи на прием для ветеринара', level=1)

        for appt in self.appointments:
            vet_document.add_paragraph(
                f'Запись на прием для {appt.animal.name}, Клиент: {appt.client.name}, '
                f'Телефон: {appt.client.phone}, Email: {appt.client.email}, '
                f'Дата приема: {appt.date}, Время приема: {appt.time}'
            )

        vet_document.save(f'{vet_directory}/{veterinarian_name}_{filename}_{current_time}.docx')

    def get_current_date(self):
        return datetime.now().strftime("%Y-%m-%d")